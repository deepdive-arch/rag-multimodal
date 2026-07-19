"""Deterministic extraction and validation for supported media."""

import wave
from dataclasses import dataclass
from pathlib import Path

import cv2
import fitz
from docx import Document
from PIL import Image, ImageSequence
from mutagen import File as MutagenFile

from core.config import Settings, get_settings
from core.exceptions import InvalidMediaError, MediaDurationExceededError
from services.chunking import ChunkDraft, TextBlock, chunk_blocks, chunk_text, normalize_text
from services.storage import file_type_for, mime_for, sanitize_filename


@dataclass(frozen=True)
class MediaItem:
    """One indexable unit produced from a file."""

    text: str
    file_type: str
    content_modality: str
    page_number: int = 0
    media_path: Path | None = None
    mime_type: str = ""
    duration_seconds: float = 0.0
    warnings: tuple[str, ...] = ()


def extract_items(path: Path, doc_id: str, settings: Settings | None = None) -> list[MediaItem]:
    """Process a supported file into text or media index items."""
    settings = settings or get_settings()
    file_type = file_type_for(path)
    if file_type == "text":
        return _text_items(path, settings)
    if file_type == "docx":
        return _docx_items(path, settings)
    if file_type == "pdf":
        return _pdf_items(path, doc_id, settings)
    if file_type == "image":
        return [_image_item(path, doc_id, settings)]
    if file_type == "audio":
        return [_audio_item(path, settings)]
    return [_video_item(path, settings)]


def _text_items(path: Path, settings: Settings) -> list[MediaItem]:
    """Extract UTF-8 text and chunk it."""
    text = path.read_text(encoding="utf-8")
    drafts = _markdown_drafts(text, settings) if path.suffix.lower() == ".md" else chunk_text(text, settings=settings)
    return [_draft_item(draft, "text", settings) for draft in drafts]


def _markdown_drafts(text: str, settings: Settings) -> list[ChunkDraft]:
    """Chunk Markdown while carrying the latest heading."""
    heading = ""
    blocks: list[TextBlock] = []
    for block in normalize_text(text).split("\n\n"):
        lines = block.splitlines()
        if lines and lines[0].lstrip().startswith("#"):
            heading = lines[0].strip()
        content = "\n".join(lines[1:]).strip() if lines and lines[0].lstrip().startswith("#") else block.strip()
        if content:
            blocks.append(TextBlock(content, 0, heading))
    return chunk_blocks(blocks, settings)


def _docx_items(path: Path, settings: Settings) -> list[MediaItem]:
    """Extract DOCX paragraphs, headings and table rows in document order."""
    document = Document(path)
    blocks: list[TextBlock] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            heading = text if paragraph.style.name.lower().startswith("heading") else ""
            blocks.append(TextBlock(text, 0, heading))
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                blocks.append(TextBlock(" | ".join(f"Coluna {index + 1}: {value}" for index, value in enumerate(values) if value), 0, "Tabela"))
    return [_draft_item(draft, "docx", settings) for draft in chunk_blocks(blocks, settings)]


def _pdf_items(path: Path, doc_id: str, settings: Settings) -> list[MediaItem]:
    """Extract PDF text or persist a rendered image for scanned pages."""
    document = fitz.open(path)
    try:
        if document.needs_pass:
            raise InvalidMediaError("PDF protegido por senha não pode ser processado")
        if document.page_count > settings.max_pdf_pages:
            raise InvalidMediaError(f"PDF excede o limite de {settings.max_pdf_pages} páginas")
        items: list[MediaItem] = []
        for index, page in enumerate(document):
            if page.rect.width * page.rect.height * (150 / 72) ** 2 > settings.max_pdf_page_pixels:
                raise InvalidMediaError("PDF contém uma página com geometria excessiva")
            text = normalize_text(page.get_text("text"))
            if len(text) >= 50:
                items.extend(_draft_item(draft, "pdf", settings) for draft in chunk_text(text, index + 1, settings=settings))
                continue
            items.append(_rendered_page_item(page, doc_id, index + 1, settings))
        return items
    finally:
        document.close()


def _rendered_page_item(page: fitz.Page, doc_id: str, page_number: int, settings: Settings) -> MediaItem:
    """Render and persist one scanned PDF page."""
    output_dir = settings.temp_processing_dir / sanitize_filename(doc_id) / "derived"
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"page-{page_number:04d}.png"
    pixmap = page.get_pixmap(dpi=150, alpha=False)
    pixmap.save(output)
    return MediaItem("", "pdf", "image", page_number, output, "image/png")


def _image_item(path: Path, doc_id: str, settings: Settings) -> MediaItem:
    """Validate an image and normalize GIF/WebP when needed."""
    try:
        with Image.open(path) as image:
            image.verify()
        with Image.open(path) as image:
            if image.width * image.height > settings.max_image_pixels:
                raise InvalidMediaError("Imagem excede o limite de pixels")
            if path.suffix.lower() == ".gif" or path.suffix.lower() == ".webp":
                output_dir = settings.temp_processing_dir / sanitize_filename(doc_id) / "derived"
                output_dir.mkdir(parents=True, exist_ok=True)
                output = output_dir / "normalized-image.jpg"
                frame = next(ImageSequence.Iterator(image)).convert("RGB")
                frame.save(output, format="JPEG", quality=92)
                warning = "GIF animado processado usando apenas o primeiro frame." if getattr(image, "is_animated", False) else "Imagem WebP normalizada para JPEG."
                return MediaItem("", "image", "image", media_path=output, mime_type="image/jpeg", warnings=(warning,))
    except InvalidMediaError:
        raise
    except Exception as error:
        raise InvalidMediaError("Imagem corrompida ou ilegível") from error
    return MediaItem("", "image", "image", media_path=path, mime_type=mime_for(path))


def _audio_item(path: Path, settings: Settings) -> MediaItem:
    """Validate audio and enforce duration limits."""
    duration = 0.0
    try:
        if path.suffix.lower() == ".wav":
            with wave.open(str(path), "rb") as audio:
                duration = audio.getnframes() / max(audio.getframerate(), 1)
        else:
            audio = MutagenFile(path)
            duration = float(audio.info.length) if audio and audio.info else 0.0
    except Exception as error:
        raise InvalidMediaError("Áudio corrompido ou ilegível") from error
    if duration <= 0:
        raise InvalidMediaError("Não foi possível determinar a duração do áudio")
    if duration > settings.max_audio_duration_seconds:
        raise MediaDurationExceededError("Áudio excede a duração máxima configurada")
    return MediaItem("", "audio", "audio", media_path=path, mime_type=mime_for(path), duration_seconds=duration)


def _video_item(path: Path, settings: Settings) -> MediaItem:
    """Validate video and calculate duration with OpenCV."""
    capture = cv2.VideoCapture(str(path))
    try:
        fps = float(capture.get(cv2.CAP_PROP_FPS))
        frames = float(capture.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frames / fps if fps > 0 else 0.0
        if not capture.isOpened() or duration <= 0:
            raise InvalidMediaError("Vídeo corrompido ou ilegível")
    finally:
        capture.release()
    if duration > settings.max_video_duration_seconds:
        raise MediaDurationExceededError("Vídeo excede a duração máxima configurada")
    warning = "Neste MVP, a faixa de áudio interna do vídeo não é indexada separadamente."
    return MediaItem("", "video", "video", media_path=path, mime_type=mime_for(path), duration_seconds=duration, warnings=(warning,))


def _draft_item(draft: ChunkDraft, file_type: str, settings: Settings) -> MediaItem:
    """Convert a chunk draft into a text media item."""
    return MediaItem(draft.text, file_type, "text", draft.page_number, mime_type="text/plain")
