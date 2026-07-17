from pathlib import Path

import fitz
from docx import Document
from PIL import Image

from core.config import Settings
from services.media import extract_items


def make_settings(tmp_path: Path) -> Settings:
    return Settings(uploads_dir=tmp_path / "uploads", derived_dir=tmp_path / "derived", database_path=tmp_path / "db.sqlite")


def test_docx_paragraphs_and_tables(tmp_path: Path):
    path = tmp_path / "document.docx"
    document = Document()
    document.add_heading("Título", level=1)
    document.add_paragraph("Conteúdo importante")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    document.save(path)
    items = extract_items(path, "doc", make_settings(tmp_path))
    joined = " ".join(item.text for item in items)
    assert "Conteúdo importante" in joined
    assert "Coluna 1: A" in joined


def test_scanned_pdf_persists_page_png(tmp_path: Path):
    settings = make_settings(tmp_path)
    path = tmp_path / "scan.pdf"
    document = fitz.open()
    document.new_page()
    document.save(path)
    document.close()
    items = extract_items(path, "doc", settings)
    assert items[0].content_modality == "image"
    assert (settings.derived_dir / "doc" / "page-0001.png").exists()


def test_webp_is_normalized(tmp_path: Path):
    settings = make_settings(tmp_path)
    path = tmp_path / "image.webp"
    Image.new("RGB", (10, 10), "red").save(path, "WEBP")
    item = extract_items(path, "doc", settings)[0]
    assert item.mime_type == "image/jpeg"
    assert item.media_path and item.media_path.suffix == ".jpg"
